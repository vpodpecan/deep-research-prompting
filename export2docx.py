import argparse
import json
import os
import re
import sqlite3
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document

# ============================================================
#                   JSON extraction
# ============================================================

def extract_answer_markdown(raw_json):
    """
    Extract the final Markdown answer from Deep Research JSON.
    """
    body = raw_json.get("body", {})
    output = body.get("output", [])

    for item in output:
        if item.get("type") == "message":
            for c in item.get("content", []):
                if c.get("type") == "output_text":
                    return c.get("text", "")
    return ""


# ============================================================
#               Word hyperlink helper
# ============================================================

def add_hyperlink(paragraph, text, url):
    """
    Creates a clickable hyperlink in a python-docx paragraph.

    paragraph: a docx.paragraph.Paragraph object
    text: the text shown to the user
    url: the actual link destination
    """
    # 1) create relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # 2) create w:hyperlink element with that r:id
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 3) create a w:r with w:rPr (for styling)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # color: blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    new_run.append(rPr)

    # 4) add the text element
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    # 5) glue it all together
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return paragraph


# ============================================================
#       Minimal Markdown → DOCX converter (with hyperlinks)
# ============================================================

# Regex for Markdown links: [text](url)
MD_LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def add_markdown_paragraph(doc, md_text):
    """
    Adds a Markdown-formatted line into the document:

      - #, ##, ### headings
      - **bold**, *italic*
      - [text](url) as REAL Word hyperlinks
    """
    md_text = md_text.rstrip()

    # ---- HEADINGS ----
    if md_text.startswith("### "):
        doc.add_heading(md_text[4:], level=3)
        return
    elif md_text.startswith("## "):
        doc.add_heading(md_text[3:], level=2)
        return
    elif md_text.startswith("# "):
        doc.add_heading(md_text[2:], level=1)
        return

    # ---- NORMAL PARAGRAPH ----
    p = doc.add_paragraph()

    pos = 0
    while True:
        match = MD_LINK_RE.search(md_text, pos)
        if not match:
            text_part = md_text[pos:]
            add_text_with_styles(p, text_part)
            break

        before = md_text[pos:match.start()]
        add_text_with_styles(p, before)

        link_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(p, link_text, link_url)

        pos = match.end()


def add_text_with_styles(paragraph, text):
    """
    Adds plain text with **bold** and *italic* markdown support.
    """
    bold_re = re.compile(r'\*\*(.*?)\*\*')
    italic_re = re.compile(r'\*(.*?)\*')

    pos = 0
    while pos < len(text):
        b = bold_re.search(text, pos)
        i = italic_re.search(text, pos)

        matches = [m for m in [b, i] if m]
        if not matches:
            paragraph.add_run(text[pos:])
            break

        m = min(matches, key=lambda m: m.start())

        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        content = m.group(1)
        run = paragraph.add_run(content)

        if m.re == bold_re:
            run.bold = True
        else:
            run.italic = True

        pos = m.end()


# ============================================================
#                 DOCX creation
# ============================================================

def write_markdown_to_word(filename, md_text):
    doc = Document()
    doc.add_heading("Deep Research Answer", level=1)

    for line in md_text.split("\n"):
        if line.strip():
            add_markdown_paragraph(doc, line)
        else:
            doc.add_paragraph("")

    doc.save(filename)


# ============================================================
#              DB export logic
# ============================================================

def safe_filename(base: str, max_len: int = 80) -> str:
    base = base or "response"
    base = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in base)
    return base[:max_len] or "response"


def export_db_to_docx(db_path: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute("SELECT id, custom_id, response, error FROM responses")
    rows = cur.fetchall()

    exported = 0
    skipped = 0

    for row in rows:
        raw = row["response"]
        error = row["error"]

        if error or not raw:
            skipped += 1
            continue

        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            print(f"Row id={row['id']}: invalid JSON, skipped.")
            skipped += 1
            continue

        md_text = extract_answer_markdown(data)
        if not md_text:
            print(f"Row id={row['id']}: no Markdown answer, skipped.")
            skipped += 1
            continue

        base = safe_filename(row["custom_id"] or f"response_{row['id']}")
        filename = os.path.join(output_dir, f"{base}.docx")

        write_markdown_to_word(filename, md_text)
        print(f"Exported: {filename}")

        exported += 1

    print(f"\nDone. Exported {exported} responses, skipped {skipped}.")


# ============================================================
#                        CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Export Deep Research Markdown answers from SQLite DB to Word with real hyperlinks."
    )
    parser.add_argument("db_path", help="Path to SQLite DB")
    parser.add_argument("--output-dir", default="docx_export",
                        help="Output directory (default: docx_export)")

    args = parser.parse_args()
    export_db_to_docx(args.db_path, args.output_dir)


if __name__ == "__main__":
    main()
