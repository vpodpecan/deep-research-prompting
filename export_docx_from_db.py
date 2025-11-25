#!/usr/bin/env python
import argparse
import json
import os
import re
import sqlite3
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import sys
import unicodedata


# Invalid XML 1.0 characters
INVALID_XML_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F"
    r"\ud800-\udfff]"
)

CONTROL_CHAR_NAMES = {
    0x00: "NULL",
    0x01: "SOH",
    0x02: "STX",
    0x03: "ETX",
    0x04: "EOT",
    0x05: "ENQ",
    0x06: "ACK",
    0x07: "BEL",
    0x08: "BS",
    0x0B: "VT",
    0x0C: "FF",
    0x0E: "SO",
    0x0F: "SI",
    0x10: "DLE",
    0x11: "DC1",
    0x12: "DC2",
    0x13: "DC3",
    0x14: "DC4",
    0x15: "NAK",
    0x16: "SYN",
    0x17: "ETB",
    0x18: "CAN",
    0x19: "EM",
    0x1A: "SUB",
    0x1B: "ESC",
    0x1C: "FS",
    0x1D: "GS",
    0x1E: "RS",
    0x1F: "US",
    0x7F: "DEL"
}

def describe_char(c):
    """Human-readable description of a control character."""
    code = ord(c)
    return CONTROL_CHAR_NAMES.get(code, unicodedata.name(c, "UNKNOWN"))


def sanitize_xml_string(s: str, context_label: str = "") -> str:
    """
    Remove invalid XML characters and log what was removed.
    context_label helps identify which part of text is being processed.
    """
    if not isinstance(s, str):
        return s

    cleaned = []
    for i, c in enumerate(s):
        if INVALID_XML_RE.match(c):
            name = describe_char(c)
            sys.stderr.write(
                f"[XML-SANITIZER] Removed U+{ord(c):04X} '{c}' ({name}) "
                f"at index {i} in {context_label}\n"
            )
            continue
        cleaned.append(c)

    return "".join(cleaned)

# ============================================================
#                     JSON Extraction
# ============================================================

def extract_markdown_answer(raw_json):
    """
    Extract Markdown answer text from Deep Research output JSON.
    """
    body = raw_json.get("body", {})
    output = body.get("output", [])

    for item in output:
        if item.get("type") == "message":
            for c in item.get("content", []):
                if c.get("type") == "output_text":
                    return c.get("text", "")

    return ""


# ============================================================
#            Word Hyperlink Creation (python-docx)
# ============================================================

def add_hyperlink(paragraph, text, url):
    """
    Insert a clickable Word hyperlink into a paragraph.
    """
    text = sanitize_xml_string(text, context_label="hyperlink_text")
    url = sanitize_xml_string(url, context_label="hyperlink_url")

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create <w:r> run container
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ============================================================
#          Markdown → Word Conversion
# ============================================================

LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def add_markdown_line(doc, line):
    line = sanitize_xml_string(line, context_label="parsed_markdown_line")
    line = line.rstrip()

    # ---- HEADINGS ----
    if line.startswith("### "):
        doc.add_heading(line[4:], level=3)
        return
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)
        return
    elif line.startswith("# "):
        doc.add_heading(line[2:], level=1)
        return

    # ---- NORMAL PARAGRAPHS ----
    p = doc.add_paragraph()

    pos = 0
    while True:
        m = LINK_RE.search(line, pos)
        if not m:
            add_styled_text(p, line[pos:])
            return

        before = line[pos:m.start()]
        add_styled_text(p, before)

        link_text = m.group(1)
        link_url = m.group(2)
        add_hyperlink(p, link_text, link_url)

        pos = m.end()


def add_styled_text(paragraph, text):
    """
    Add normal + styled (bold/italic) text into paragraph.
    """
    text = sanitize_xml_string(text, context_label="styled_text")

    bold_re = re.compile(r'\*\*(.*?)\*\*')
    italic_re = re.compile(r'\*(.*?)\*')

    pos = 0
    while pos < len(text):
        b = bold_re.search(text, pos)
        i = italic_re.search(text, pos)

        matches = [m for m in (b, i) if m]
        if not matches:
            paragraph.add_run(text[pos:])
            return

        m = min(matches, key=lambda x: x.start())

        # Add preceding plain text
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        content = m.group(1)
        run = paragraph.add_run(content)

        if m.re == bold_re:
            run.bold = True
        elif m.re == italic_re:
            run.italic = True

        pos = m.end()


# ============================================================
#                       DOCX Writer
# ============================================================

def write_docx(filename, markdown_text):
    doc = Document()
    doc.add_heading("Deep Research Answer", level=1)

    for line in markdown_text.split("\n"):
        line = sanitize_xml_string(line, context_label="markdown_line")
        if line.strip():
            add_markdown_line(doc, line)
        else:
            doc.add_paragraph("")

    doc.save(filename)


# ============================================================
#                         Exporter
# ============================================================

def safe_filename(s: str, max_len: int = 80) -> str:
    if not s:
        return "response"
    s = "".join(c if c.isalnum() or c in "-_" else "_" for c in s)
    return s[:max_len] or "response"


def export_to_word(db_path: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute("""
        SELECT batch_id, custom_id, prompt, response, error
        FROM responses
        WHERE response IS NOT NULL
    """)
    rows = cur.fetchall()

    exported = 0
    skipped = 0

    for row in rows:
        b_id = row["batch_id"]
        custom_id = row["custom_id"]
        response_json = row["response"]
        error = row["error"]

        if error:
            skipped += 1
            continue

        if not response_json:
            skipped += 1
            continue

        try:
            data = json.loads(response_json)
        except json.JSONDecodeError:
            skipped += 1
            continue

        markdown = extract_markdown_answer(data)
        if not markdown:
            skipped += 1
            continue

        #fname = safe_filename(custom_id) + ".docx"
        fname = f'{custom_id}' + ".docx"
        path = os.path.join(output_dir, fname)

        write_docx(path, markdown)
        print(f"Exported: {path}")
        exported += 1

    print(f"\nDone. Exported {exported} documents, skipped {skipped}.")


# ============================================================
#                         CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Export Deep Research Markdown answers (with hyperlinks) from SQLite to Word."
    )
    parser.add_argument("db_path", help="SQLite DB containing batches + responses tables")
    parser.add_argument("--output-dir", default="docx_export",
                        help="Directory to store .docx files")

    args = parser.parse_args()
    export_to_word(args.db_path, args.output_dir)


if __name__ == "__main__":
    main()
