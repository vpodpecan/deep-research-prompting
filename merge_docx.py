#!/usr/bin/env python3
import argparse
import os
import sys
from docx import Document
from docxcompose.composer import Composer


def set_heading_style(p):
    """Try common heading style names across Word locales/templates."""
    for name in ("Heading 1", "Heading1"):
        try:
            p.style = name
            return
        except Exception:
            pass
    # fallback: leave default style if heading style name not found


def prepend_filename_heading(doc: Document, filename: str):
    """
    Insert filename heading at the very beginning of the doc without modifying original file on disk.
    """
    if doc.paragraphs:
        first = doc.paragraphs[0]
        p = first.insert_paragraph_before(filename)
    else:
        p = doc.add_paragraph(filename)

    set_heading_style(p)


def merge(files, output):
    # keep only existing docx
    files = [f for f in files if os.path.isfile(f) and f.lower().endswith(".docx")]
    if not files:
        print("No .docx files found.", file=sys.stderr)
        sys.exit(1)

    # sort alphabetically (case-insensitive)
    files = sorted(files, key=lambda x: os.path.basename(x).lower())

    print("Merging documents in this order:")
    for f in files:
        print(" -", f)

    # Load first doc as master
    master_path = files[0]
    master_doc = Document(master_path)
    prepend_filename_heading(master_doc, os.path.basename(master_path))
    # Ensure the first document ends with a page break if more docs follow
    if len(files) > 1:
        master_doc.add_page_break()

    composer = Composer(master_doc)

    # Append remaining docs
    for i, path in enumerate(files[1:], start=1):
        d = Document(path)

        # Put filename heading on its own page at top of this appended doc
        prepend_filename_heading(d, os.path.basename(path))

        # Add a page break at the end of this doc unless it's the last one
        if i < len(files) - 1:
            d.add_page_break()

        composer.append(d)

    composer.save(output)
    print(f"\nMerged document saved as: {output}")


def main():
    ap = argparse.ArgumentParser(
        description="Merge multiple DOCX files into one, each starting on a new page with filename heading (hyperlinks preserved)."
    )
    ap.add_argument("files", nargs="+", help="Input .docx files (shell glob like ABC*.docx is fine).")
    ap.add_argument("-o", "--output", default="merged.docx", help="Output file (default: merged.docx)")
    args = ap.parse_args()

    merge(args.files, args.output)


if __name__ == "__main__":
    main()
